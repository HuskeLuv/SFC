#!/usr/bin/env python3
"""Gera o .docx executivo (para a diretoria) da avaliação da API Cedro."""
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F, 0x33, 0x5E)
GREY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
RED = RGBColor(0xB0, 0x2A, 0x2A)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def para(text="", italic=False, color=None, size=None, bold=False, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    return t


# ---- Título ----
# (edições do usuário em 02/07 preservadas: sem subtítulo/data/"Em uma frase";
#  seção 1 sem o ", sem jargão")
title = doc.add_heading("Avaliação da API Cedro", level=0)
for run in title.runs:
    run.font.color.rgb = NAVY

# ---- 1 ----
heading("1. O que está em jogo", 1)
para(
    "Nosso aplicativo precisa de “dados de mercado” para funcionar: quanto vale "
    "cada ativo hoje, quanto o cliente recebeu de dividendos, e como o preço se comportou "
    "ao longo dos anos para desenhar os gráficos de rentabilidade. Hoje buscamos esses "
    "dados em várias fontes gratuitas ou de baixo custo e “costuramos” elas juntas. "
    "Isso funciona, mas dá trabalho e, de vez em quando, gera pequenas divergências (foi a "
    "causa de vários ajustes que fizemos nos últimos meses)."
)
para(
    "A Cedro é um fornecedor profissional de dados da Bolsa (B3). A promessa é: em vez de "
    "costurar cinco fontes, ter uma fonte só, mais completa e confiável. Recebemos um acesso "
    "de teste (“trial”) e passamos os últimos dias comparando os dados dela com os "
    "que já usamos. Este documento é o resultado."
)

# ---- 2 ----
heading("2. O que usamos hoje (as fontes atuais)", 1)
table(
    ["Fonte", "Para quê usamos", "Custo"],
    [
        ["BRAPI", "Cotação do dia e dividendos recentes — fonte principal", "Assinatura paga (plano fixo)"],
        ["Yahoo Finance", "Splits/desdobramentos e histórico antigo", "Gratuita"],
        ["Arquivo oficial da B3 (COTAHIST)", "Histórico de preço oficial desde 2016", "Gratuita (arquivo público)"],
        ["Tesouro / BACEN / CVM", "Títulos públicos, índices e fundos", "Gratuitas (oficiais)"],
        ["CoinGecko", "Criptomoedas", "Gratuita"],
    ],
)
para("", space_after=2)
para(
    "Ou seja: já temos uma base sólida e barata. A pergunta não é “temos dados?”, e "
    "sim: a Cedro melhora a qualidade a ponto de justificar o custo?",
    bold=True,
)

# ---- 3 ----
heading("3. Comparação, quesito por quesito", 1)
table(
    ["Quesito", "Como estamos hoje", "O que a Cedro entrega", "Veredito"],
    [
        ["Dividendos, JCP e rendimentos de FII",
         "Combinamos BRAPI + Yahoo; às vezes falta a “data-com”",
         "Tudo de uma fonte só (valor, tipo e as 3 datas); mais completo para FII",
         "Cedro ganha"],
        ["Splits / desdobramentos",
         "Yahoo é a base; já causou bug de duplicação",
         "Classifica o evento com mais precisão, mas deixou passar 1 split antigo de FII",
         "Empate / usar as duas"],
        ["Histórico de preço (gráficos)",
         "Arquivo oficial da B3 desde 2016 + BRAPI",
         "Série longa, mas em formato que não encaixa direto; nossa fonte já é oficial",
         "Não substitui hoje"],
        ["Cotação do dia",
         "BRAPI já resolve",
         "Extras (compra/venda, máx/mín) que hoje não temos",
         "Complementa"],
        ["Lista de ativos da Bolsa",
         "BRAPI + CVM",
         "Redundante com o que já temos",
         "Só conferência"],
    ],
)
para("", space_after=2)
para(
    "Leitura rápida: a Cedro é claramente melhor em proventos, é uma boa segunda opinião em "
    "splits, e não muda o jogo no resto.",
    bold=True,
)

# ---- 4 ----
heading("4. O exemplo do split — HFOF11 (o teste mais concreto)", 1)
para(
    "Pedimos para focar num ativo que passou por um split (desdobramento), porque é "
    "justamente onde os dados costumam divergir e onde já tivemos bugs. Escolhemos o fundo "
    "imobiliário HFOF11, que fez um desdobramento de 10 para 1 em maio de 2025 — cada cota "
    "virou 10 cotas, e o preço foi naturalmente dividido por 10."
)

heading("4.1 O evento em si — quem reporta certo?", 2)
table(
    ["Fonte", "Reporta o split?", "Fator", "Datas", "Observação"],
    [
        ["Cedro", "Sim, 1 evento único", "10:1 (“+900%”)", "com 09/05, ex 12/05, pag 13/05",
         "Classifica como “Desdobramento”; traz as 3 datas"],
        ["Yahoo Finance", "Sim", "10:1", "ex 12/05", "Correto, mas sem a data-com"],
        ["BRAPI", "Não traz o split isolado", "—", "—", "Por isso combinávamos com o Yahoo"],
        ["Nosso sistema (antes do ajuste)", "Sim, porém DUPLICADO", "20:1 (errado!)", "—",
         "BRAPI + Yahoo somavam o mesmo evento — bug que corrigimos"],
    ],
)
para("", space_after=2)
para(
    "Conclusão: a Cedro, por ser fonte única, elimina na raiz o risco de duplicação que nos "
    "causou dor de cabeça — e entrega a “data-com”, que o Yahoo não tem.",
    bold=True,
)

heading("4.2 Por que isso importa — a “queda-fantasma”", 2)
para(
    "Este é o ponto que vale a diretoria entender, porque é onde o cliente sente no gráfico. "
    "Veja o preço do HFOF11 ao redor do split:"
)
table(
    ["Dia", "Preço SEM tratamento (nominal)", "Preço TRATADO (correto)"],
    [
        ["09/05/2025 (antes)", "≈ R$ 52,90", "R$ 5,29"],
        ["12/05/2025 (dia do split)", "R$ 5,26", "R$ 5,26"],
        ["Variação aparente", "− 90%", "estável"],
    ],
)
para("", space_after=2)
para(
    "Sem o tratamento correto do split, o gráfico do cliente mostraria uma queda de 90% da "
    "noite para o dia — uma perda que nunca existiu (ele passou a ter 10× mais cotas valendo "
    "1/10 cada). Tanto a Cedro quanto nossa correção atual já fazem isso corretamente; a "
    "diferença é que com a Cedro viria pronto e de uma fonte só, em vez de costurado."
)
para(
    "Detalhe que confirma o evento: no dia do split o volume negociado saltou de ~17 mil para "
    "~204 mil cotas — comportamento típico de um desdobramento, e que bate em todas as fontes.",
    italic=True, color=GREY, size=10,
)

heading("4.3 A prova no gráfico — rentabilidade dia a dia: nossos dados × Cedro × Kinvo", 2)
para(
    "Para fechar a comparação, colocamos lado a lado a rentabilidade diária acumulada do "
    "HFOF11 nos ~4 meses ao redor do split (mar–jun/2025), de três fontes independentes: "
    "nossos dados de produção (arquivo oficial da B3), a série da Cedro (captura ao vivo de "
    "30/06) e o Kinvo — o app concorrente que usamos como benchmark — capturado do "
    "aplicativo real em 02/07:"
)
CHART = os.path.join(os.path.dirname(__file__), "..", "docs", "cedro-captures",
                     "hfof-rentabilidade-split-jul2026.png")
doc.add_picture(CHART, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
bullet(
    "As três curvas contam a mesma história. Dia a dia, a diferença média entre qualquer "
    "par de fontes fica entre 0,10 e 0,17 ponto percentual — e nenhuma das três “sente” o "
    "split: no dia 12/05 as curvas seguem em linha reta, como deve ser (sem a queda-fantasma "
    "de −90% descrita no item 4.2)."
)
bullet(
    "Kinvo e Cedro fecham quase juntas (+14,3% × +14,0%) — as duas embutem os rendimentos "
    "mensais do fundo na série (como se fossem reinvestidos). Duas fontes independentes "
    "chegando a 0,35 ponto uma da outra é uma validação cruzada forte das duas."
)
bullet(
    "A nossa curva (+11,4%) fica ~2,9 pontos abaixo, e isso não é erro: nossa série mostra "
    "só o preço da cota, sem embutir os rendimentos — que, em 4 meses, somam justamente essa "
    "diferença. É a demonstração prática do item 3: o formato da Cedro (e do Kinvo) é "
    "diferente do nosso — bom para conferência, mas não é um encaixe direto para substituir "
    "nossa série de preço."
)
para(
    "Resumo do gráfico: os três lados tratam o split corretamente e contam a mesma história; "
    "as divergências que existem são metodológicas e explicadas (com ou sem rendimentos "
    "embutidos), não erros de dados.",
    bold=True,
)

heading("4.4 Indicadores do ativo — o que cada fonte informa", 2)
para(
    "Além do gráfico, comparamos os indicadores-resumo que aparecem na página do ativo — "
    "quanto ele rende, quanto oscila e quanto acompanha a Bolsa:"
)
CHART_IND = os.path.join(os.path.dirname(__file__), "..", "docs", "cedro-captures",
                         "hfof-indicadores-jul2026.png")
doc.add_picture(CHART_IND, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
bullet(
    "Rendimento (dividend yield) e volatilidade: as três fontes contam a mesma história "
    "(10,0–10,9% e 12,9–14,1%). As pequenas diferenças vêm da janela e do preço de "
    "referência que cada um usa — e a soma dos 12 rendimentos do ano bate ao centavo entre "
    "nossos dados e a Cedro (R$ 0,692 por cota), com o último rendimento (R$ 0,06) idêntico "
    "nas três."
)
bullet(
    "Beta é o exemplo de indicador em que a metodologia muda tudo: nosso cálculo dá 0,22 e "
    "o Kinvo mostra 0,36 — e, curiosamente, dentro do próprio Kinvo existem dois números "
    "diferentes (0,36 na tela de análise e 1,64 no serviço de indicadores fundamentalistas). "
    "Moral: comparar beta entre plataformas sem conhecer a conta por trás não é confiável — "
    "divergência aqui não indica erro de dados."
)
bullet(
    "P/VP (0,85) e valor patrimonial por cota (R$ 8,02): hoje só o Kinvo informa. Nem "
    "nossos dados nem a Cedro (no trial) trazem o patrimônio dos FIIs — se quisermos exibir "
    "esses indicadores, a fonte seria outra (informes oficiais da CVM), não a Cedro."
)
para(
    "Resumo dos indicadores: nos que dependem de dados brutos (rendimento, volatilidade), "
    "todo mundo concorda; nos que dependem de metodologia (beta), até a mesma plataforma "
    "diverge de si mesma. A Cedro não traria vantagem aqui — mais um reforço de que o valor "
    "dela está nos proventos, não nos indicadores.",
    bold=True,
)

heading("4.5 Os benchmarks do nosso gráfico — IBOV, CDI, IPCA e Poupança nos três sistemas", 2)
para(
    "O gráfico de rentabilidade do nosso app compara o ativo com as referências clássicas "
    "do investidor brasileiro: IBOV, CDI, IPCA e Poupança. Verificamos, uma a uma, quem tem "
    "cada série e se os números batem — nossos dados × Cedro (consulta ao vivo de 02/07) × "
    "Kinvo (captura do app real):"
)
CHART_BENCH = os.path.join(os.path.dirname(__file__), "..", "docs", "cedro-captures",
                           "hfof-benchmarks-jul2026.png")
doc.add_picture(CHART_BENCH, width=Inches(6.3))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
bullet(
    "IBOV — três fontes, praticamente idênticas: +12,85% (nós) × +12,85% (Cedro) × "
    "+12,84% (Kinvo). As três curvas ficam sobrepostas no gráfico."
)
bullet(
    "CDI — nossos dados × Kinvo: 0,01 ponto de diferença (+4,27% × +4,28%). Visualmente "
    "indistinguíveis."
)
bullet("IPCA — nossos dados × Kinvo: idênticos (+1,47% nos dois).")
bullet(
    "Poupança — +2,53% (nós) × +2,61% (Kinvo). Diferença de 0,08 ponto em 4 meses, "
    "explicada pela convenção de cálculo (a poupança rende na “data de aniversário” do "
    "depósito; cada sistema dilui isso de um jeito)."
)
bullet(
    "Cedro: no nosso acesso de teste, só o IBOV está liberado. Os mercados de índices/"
    "indicadores econômicos existem no protocolo dela, mas o trial responde “sem "
    "permissão” (erro E:3) — então não dá para avaliar CDI/IPCA/Poupança da Cedro sem "
    "contratar (ou pedir liberação no trial)."
)
para(
    "Resumo dos benchmarks: nós e o Kinvo cobrimos as quatro referências — e tudo o que se "
    "sobrepõe, bate (a maior diferença foi 0,08 ponto na Poupança, por convenção de "
    "cálculo). A Cedro, no acesso atual, só mostra o IBOV — que também bate. Nossos dados "
    "de referência (BACEN + B3) estão validados contra os dois concorrentes.",
    bold=True,
)

heading("4.6 A honestidade da avaliação — onde a Cedro falhou", 2)
para(
    "Para não passar uma imagem só positiva: testamos também o MXRF11, outro FII que fez um "
    "split de 10:1 em 2017. A Cedro não reportou esse evento — mas o arquivo oficial da B3 "
    "confirma que ele é real (o preço saiu de R$ 89 para R$ 9,98). Ou seja, a Cedro não é "
    "infalível para eventos antigos de FII. Por isso a recomendação em splits é usar Cedro e "
    "Yahoo em conjunto, e não trocar uma pela outra."
)
para("Essa falha foi verificada com rigor, por quatro caminhos independentes:")
bullet(
    "1. Nosso banco de produção (arquivo oficial da B3): o preço cai de R$ 89 para R$ 9,98 "
    "na data — o split é real;"
)
bullet(
    "2. Fontes públicas independentes (Investing.com, Suno, XP, StatusInvest): todas "
    "registram o desdobramento de 10:1 em 17/05/2017;"
)
bullet(
    "3. Reconferência ao vivo na própria Cedro (02/07): consultamos de novo, em duas janelas "
    "de datas — o evento realmente não é reportado;"
)
bullet(
    "4. A “impressão digital” nos dados da própria Cedro: os rendimentos mensais que ela "
    "informa caem 10 vezes (de ~R$ 0,95 para ~R$ 0,10 por cota) exatamente na data do split "
    "— ou seja, a base dela reflete o evento, mas não o expõe."
)

# ---- 5 ----
heading("5. Pontos comerciais a resolver antes de fechar", 1)
numbered(
    "Custo: ainda não temos o modelo de cobrança da Cedro (valor fixo? por ativo/consulta/"
    "usuário?). Hoje a BRAPI é um valor fixo previsível — precisamos comparar."
)
numbered(
    "Licença de redistribuição da B3: repassar cotação da Bolsa ao cliente final exige "
    "licença de dados de mercado. Confirmar se já está incluída no contrato."
)
numbered(
    "Prazo: o acesso de teste expira em 03/07/2026. A avaliação técnica já está concluída; a "
    "decisão comercial pode seguir sem pressa do trial."
)

# ---- 6 ----
heading("6. Recomendação", 1)
bullet(
    "Contratar a Cedro especificamente para PROVENTOS (dividendos, JCP e rendimentos de FII) "
    "é o passo de maior retorno — é onde ela é claramente superior e resolve dores reais."
)
bullet("Splits: usar a Cedro somando ao que já temos, nunca substituindo.")
bullet(
    "Histórico de preço e cotação: manter o que temos — nossa fonte já é a oficial da B3 e a "
    "Cedro não justifica a troca no plano atual."
)
bullet("Antes de assinar: fechar as pendências comerciais do item 5 (custo e licença B3).")
para("", space_after=2)
para(
    "Em resumo: é um bom fornecedor para um problema específico e valioso (proventos), mas "
    "não um substituto de tudo. A adoção recomendada é cirúrgica, não total — o que também "
    "limita o custo.",
    bold=True, color=NAVY,
)

import sys
out = sys.argv[1] if len(sys.argv) > 1 else "Relatorio-Cedro-Diretoria-jul2026.docx"
doc.save(out)
print("Saved:", out)
