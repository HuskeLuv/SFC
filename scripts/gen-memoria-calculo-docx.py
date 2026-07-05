#!/usr/bin/env python3
"""Gera o .docx da Memória de Cálculo do gráfico de rentabilidade diária a partir do MD."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x1F, 0x33, 0x5E)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h

def para(text="", italic=False, color=None, size=None, bold=False, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")

def callout(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run("▶  " + text)
    r.italic = True
    r.font.color.rgb = GREY
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(8)
    return p

def formula(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = NAVY
    return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

# ---------- Capa / título ----------
title = doc.add_heading("Memória de Cálculo", level=0)
for run in title.runs:
    run.font.color.rgb = NAVY
para("Gráfico de Rentabilidade Diária — Análise da Carteira", bold=True, size=13)
para("MyFinance · Documento explicativo (linguagem de negócio) · Atualizado em 15/06/2026",
     italic=True, color=GREY, size=10, space_after=14)

para("O que este documento explica: de onde vem cada número do gráfico de rentabilidade "
     "diária da seção de Análises e por que ele é calculado daquele jeito. Cobre a linha da "
     "própria carteira e as quatro linhas de comparação (IBOV, CDI, IPCA e Poupança). "
     "Linguagem de negócio, sem código.", italic=True, color=GREY, size=10, space_after=14)

# ---------- 1 ----------
heading("1. O que o gráfico mostra", 1)
para("Todas as linhas do gráfico representam rentabilidade acumulada em %, começando em 0% "
     "no primeiro dia do período e subindo ou descendo conforme o desempenho.")
bullet("A linha da carteira mostra quanto o seu dinheiro rendeu (preço dos ativos + renda "
       "recebida), neutralizando o efeito de quando você aportou ou resgatou.")
bullet("As linhas de IBOV, CDI, IPCA e Poupança são réguas de comparação: “se eu tivesse "
       "seguido esse índice no mesmo período, teria rendido X%”.")
para("Por começarem todas em zero no mesmo dia, é possível comparar visualmente “ganhei do "
     "CDI?”, “acompanhei o IBOV?” etc.")

# ---------- 2 ----------
heading("2. A linha da carteira", 1)

heading("2.1 A ideia central: retorno total, ajustado pelo tempo", 2)
para("A carteira é medida por TWR (Time-Weighted Return — retorno ponderado pelo tempo). É a "
     "mesma metodologia usada pelo Kinvo e pela indústria de fundos.")
para("O TWR responde à pergunta: “quão bem os meus investimentos performaram?” — e não "
     "“quanto dinheiro eu tenho a mais”. A diferença importa:")
bullet("Se você aporta um valor grande hoje e a carteira cai amanhã, o TWR quase não se mexe, "
       "porque o dinheiro novo ainda não teve tempo de render.")
bullet("O TWR isola o desempenho dos ativos do tamanho e do momento dos seus aportes.")
para("Isso é o que permite comparar de forma justa com IBOV/CDI: estamos comparando qualidade "
     "de desempenho, não volume de aporte.")

heading("2.2 Como cada dia é calculado, em termos simples", 2)
para("Para cada dia útil da B3, o sistema calcula o valor total da carteira (“saldo bruto”):")
para("Saldo do dia  =  valor de mercado dos ativos  +  caixa (dinheiro não investido)  +  "
     "proventos já recebidos (acumulados)", italic=True, color=GREY)
para("A partir daí, mede-se quanto a carteira rendeu de um dia para o outro, descontando o que "
     "entrou ou saiu de dinheiro novo naquele dia (aportes e resgates). Esse rendimento diário "
     "é encadeado dia após dia, formando a curva acumulada.")
callout("Por quê descontar aportes/resgates? Porque dinheiro que você depositou não é "
        "“rendimento” — é capital novo. Se não descontasse, um aporte grande pareceria, "
        "falsamente, uma valorização enorme da carteira.")

heading("2.3 De onde vem cada componente", 2)
table(["Componente", "De onde vem", "Por quê"], [
    ["Quantidade de cada ativo",
     "Reconstruída do histórico de compras e vendas, já considerando splits, grupamentos e bonificações",
     "Mantém a quantidade do passado na mesma escala de hoje, senão o gráfico daria saltos artificiais na data do split"],
    ["Preço de mercado",
     "Histórico de preços (B3 / BRAPI / Yahoo, nessa prioridade), já ajustado por splits",
     "É o valor real pelo qual o ativo era negociado naquele dia"],
    ["Caixa",
     "Aportes e resgates registrados",
     "Representa dinheiro disponível ainda não aplicado"],
    ["Proventos",
     "Histórico de dividendos/JCP/rendimentos por ativo, com ajuste manual possível",
     "Conta como renda recebida (ver 2.4)"],
    ["Renda fixa",
     "Valor corrigido pela taxa contratada (CDI, IPCA, prefixado ou PU do Tesouro)",
     "Renda fixa não tem cotação de bolsa; o valor vem da rentabilidade contratada (ver 2.5)"],
])

heading("2.4 Como os proventos entram (ponto sensível)", 2)
para("Proventos (dividendos, JCP, rendimentos de FII) entram como retorno, não como saque:")
bullet("Eles somam ao desempenho da carteira (você recebeu renda).")
bullet("Eles não reduzem o saldo nem contam como aporte.")
para("Quando o provento é creditado na curva: na data de pagamento (ajustada para o próximo "
     "pregão), não na data-ex. Essa escolha é deliberada para espelhar o Kinvo — é o que faz a "
     "renda acumulada bater praticamente centavo a centavo com a plataforma de referência.")
para("Para JCP, o valor creditado já é líquido de IRRF (imposto retido na fonte); dividendos "
     "são isentos e entram cheios.")
callout("Reinvestimento de proventos não é tratado como aporte novo. Se você usou um dividendo "
        "para comprar mais ações, isso não infla o TWR — o dinheiro já estava dentro da carteira.")

heading("2.5 Renda fixa", 2)
para("Ativos de renda fixa (CDB, LCI/LCA, Tesouro) não têm preço de bolsa diário. O valor de "
     "cada dia é calculado pela rentabilidade contratada:")
bullet("Prefixado: aplica a taxa anual proporcionalmente a cada dia útil.")
bullet("Pós-fixado (% do CDI): corrige pelo CDI do dia (taxa publicada pelo Banco Central).")
bullet("IPCA+: aplica a parte prefixada diariamente e incorpora o IPCA ao virar o mês.")
bullet("Tesouro Direto: usa o preço unitário (PU) oficial divulgado.")

heading("2.6 De onde o sistema lê tudo isso (rápido vs. reconstrução)", 2)
para("Existem dois caminhos para montar a série, e o resultado é o mesmo:")
bullet("Caminho rápido (padrão): lê fotografias diárias já calculadas e guardadas no banco "
       "(atualizadas por uma rotina automática diária). É o que acontece na maioria dos acessos.")
bullet("Caminho de reconstrução (fallback): quando não há fotografias suficientes (usuário novo, "
       "ou a rotina ainda não cobriu o período), o sistema recalcula tudo ao vivo a partir das "
       "transações, eventos, preços e proventos.")
para("O sistema decide sozinho qual usar, e cai no fallback de forma transparente quando detecta "
     "lacunas no histórico salvo.")

# ---------- 3 ----------
heading("3. As linhas de comparação (benchmarks)", 1)
para("Todas começam em 0% no mesmo dia da carteira e mostram rentabilidade acumulada, para "
     "comparação direta.")
table(["Benchmark", "O que é", "Fonte oficial", "Frequência do dado"], [
    ["IBOV", "Índice da bolsa brasileira (B3)", "Cotação do Ibovespa (Yahoo Finance / BRAPI)", "Diária"],
    ["CDI", "Referência da renda fixa pós-fixada", "Banco Central (série diária)", "Diária"],
    ["IPCA", "Inflação oficial", "Banco Central (série mensal)", "Mensal"],
    ["Poupança", "Rendimento da caderneta", "Banco Central (série mensal)", "Mensal"],
])

heading("3.1 Como cada um vira uma curva diária", 2)
bullet("IBOV: já é um número diário. A rentabilidade é quanto o índice subiu ou caiu em relação "
       "ao primeiro dia. Em fins de semana e feriados, repete-se o último valor disponível.")
bullet("CDI: o Banco Central publica a taxa de cada dia. O sistema acumula essas taxas dia a dia "
       "(“juros sobre juros”) para formar a curva.")
bullet("IPCA e Poupança: o Banco Central só publica valores mensais. O sistema acumula mês a mês "
       "e distribui de forma suave ao longo dos dias para desenhar uma linha contínua "
       "(tecnicamente, interpolação geométrica — ver Apêndice B.3).")
callout("Interpretação importante: como IPCA e Poupança são mensais “esticados” em linha reta, "
        "a oscilação diária dessas duas linhas não é dado real diário — é uma aproximação visual "
        "entre pontos mensais. A comparação justa é de tendência ao longo de semanas/meses, não "
        "do solavanco de um único dia.")

heading("3.2 De onde o sistema lê os benchmarks", 2)
para("Mesma lógica de dois caminhos da carteira, para não depender de internet a cada acesso:")
bullet("Primeiro: valores já calculados e guardados no banco (pré-carregados).")
bullet("Se faltar algum: busca direto na fonte oficial (Banco Central para CDI/IPCA/Poupança; "
       "Yahoo/BRAPI para o IBOV) e guarda para os próximos acessos.")
para("Os valores ficam em cache por até 24 horas quando os quatro benchmarks estão completos "
     "(1 hora se algum estiver faltando), para equilibrar atualidade e desempenho.")

# ---------- 4 ----------
heading("4. Resumo de uma frase por linha", 1)
bullet("Carteira: retorno total (preço dos ativos + proventos recebidos), medido por TWR para "
       "isolar desempenho do efeito dos aportes — proventos creditados na data de pagamento, "
       "igual ao Kinvo.")
bullet("IBOV: variação acumulada do índice da bolsa, dado diário real.")
bullet("CDI: juros diários do Banco Central acumulados (juros sobre juros).")
bullet("IPCA: inflação mensal oficial, acumulada e suavizada em linha diária.")
bullet("Poupança: rendimento mensal da caderneta, acumulado e suavizado em linha diária.")

# ---------- 5 ----------
heading("5. Cuidados e limitações conhecidas", 1)
bullet("IPCA e Poupança são mensais interpolados — não há precisão diária nessas curvas (ver 3.1).")
bullet("A carteira depende da qualidade dos dados de origem — preços de mercado, datas de "
       "proventos e eventos corporativos (splits). Erros de cadastro se propagam ao gráfico.")
bullet("Variações diárias muito grandes (acima de ±50% num único dia) são tratadas como ruído de "
       "dado e neutralizadas, para evitar que um erro de preço pontual distorça a curva inteira.")
bullet("Pequenas diferenças residuais vs. Kinvo podem existir por diferenças de metodologia fina "
       "(ancoragem do primeiro dia, arredondamentos), mesmo com a renda acumulada batendo.")

# ============================================================================
#  APÊNDICE TÉCNICO — VERIFICAÇÃO
# ============================================================================
doc.add_page_break()
heading("Apêndice Técnico — Verificação", 1)
para("Esta seção é o material de auditoria. Traz as fórmulas exatas, os parâmetros/convenções "
     "e uma carteira-brinquedo calculada dia a dia, para que um consultor externo reproduza "
     "os números numa planilha e confirme que o sistema está correto. Os valores foram "
     "escolhidos para fechar em números redondos, facilitando a conferência manual.",
     italic=True, color=GREY, size=10, space_after=12)

# ---- A. Parâmetros e convenções ----
heading("A. Parâmetros, fontes e convenções (exatos)", 2)
table(["Item", "Valor / Convenção"], [
    ["Calendário", "Dias úteis B3 (seg–sex, exceto feriados nacionais B3). Ancoragem em UTC."],
    ["Datas de transação", "Fim de semana/feriado é empurrado para o próximo pregão (D+próximo útil)."],
    ["Base do gráfico", "Primeiro dia do período = 0% (rentabilidade acumulada)."],
    ["Métrica da carteira", "TWR (Time-Weighted Return), encadeamento multiplicativo diário."],
    ["Clamp do retorno diário", "Se |retorno do dia| > 50% → tratado como 0 (ruído de dado)."],
    ["Clamp do dia 0", "Ganho instantâneo permitido até ±100%; fora disso → 0."],
    ["Arredondamento exibido", "2 casas decimais no %. O acumulado interno usa precisão cheia."],
    ["Renda fixa — dias/ano", "252 dias úteis (fator prefixado diário)."],
    ["Proventos — crédito", "Data de PAGAMENTO, ajustada para o próximo pregão B3."],
    ["IRRF dividendos", "0% (isento) — entra cheio."],
    ["IRRF JCP", "15% até 31/12/2025; 17,5% a partir de 01/01/2026."],
    ["IBOV", "Ticker ^BVSP (Yahoo Finance / BRAPI), cotação diária."],
    ["CDI", "BACEN SGS série 12 (taxa diária, em decimal)."],
    ["IPCA", "BACEN SGS série 433 (taxa mensal)."],
    ["Poupança", "BACEN SGS série 25 (taxa mensal)."],
])

# ---- B. Fórmulas ----
heading("B. Fórmulas exatas", 2)

heading("B.1 Saldo bruto diário (valor total da carteira)", 3)
formula("saldo_bruto(d) = Σ_i [ qtd_i(d) × preço_i(d) ]\n"
        "              + caixa(d)\n"
        "              + proventos_acumulados(d)\n"
        "              + Σ_rf [ aplicado_rf × fator_rf(d) ]")
para("qtd_i(d) já considera splits/grupamentos (ver B.6). preço_i(d) é o preço de mercado "
     "ajustado por splits. proventos_acumulados(d) é a soma líquida de todos os proventos "
     "creditados até o dia d (ver B.5). fator_rf(d) é o fator de renda fixa (ver B.4).",
     size=10, color=GREY)

heading("B.2 Retorno diário (TWR) e encadeamento", 3)
para("Para cada dia d (a partir do 2º ponto da série):")
formula("retorno(d) = ( saldo_bruto(d) − saldo_bruto(d−1) − fluxo(d) ) / saldo_bruto(d−1)\n\n"
        "se |retorno(d)| > 0,50  →  retorno(d) = 0      (clamp anti-ruído)")
para("fluxo(d) = aportes − resgates do dia d (dinheiro novo). Proventos NÃO entram em "
     "fluxo(d). Reinvestimentos de proventos NÃO entram em fluxo(d).", size=10, color=GREY)
para("Primeiro ponto da série (dia 0), quando há aporte inicial:")
formula("retorno(0) = ( saldo_bruto(0) − fluxo(0) ) / fluxo(0)     (clamp ±100%)\n"
        "Se o preço pago = preço de mercado no dia, então retorno(0) = 0.")
para("Encadeamento (produtório) e valor exibido:")
formula("acumulado(d) = Π_{k=0..d} ( 1 + retorno(k) )\n"
        "valor_no_gráfico(d) = ( acumulado(d) − 1 ) × 100   [%]")

heading("B.3 Benchmarks", 3)
para("Rebaseamento de qualquer série de índice para começar em 0% no início do período:")
formula("valor(d) = ( índice(d) / índice(base) − 1 ) × 100")
para("CDI — compõe a taxa diária do BACEN (decimal), apenas nos dias publicados:")
formula("índice(0) = 100\n"
        "índice(d) = índice(d−1) × ( 1 + taxa_cdi(d) )   [só em dia com publicação]")
para("IPCA e Poupança — índice mensal e interpolação diária geométrica:")
formula("Índice mensal:  I(mês+1) = I(mês) × ( 1 + taxa_mensal/100 ),   I(0)=100\n\n"
        "Fator diário:   f = ( I(próx. mês) / I(mês) ) ^ ( 1 / dias_do_mês )\n"
        "Valor do dia k: I(mês) × f^k")
para("Observação: a interpolação é GEOMÉTRICA (fator diário constante), não estritamente "
     "linear. Entre dois meses próximos a curva parece quase reta, mas tecnicamente é "
     "geométrica. O corpo deste documento chama isso de “suave” por simplicidade.",
     size=10, color=GREY)
para("IBOV — é cotação diária real; aplica-se diretamente o rebaseamento acima. Fins de "
     "semana/feriados repetem o último valor disponível.", size=10, color=GREY)

heading("B.4 Renda fixa — fator diário", 3)
formula("Prefixado (PRE):  fator ×= (1 + taxa_anual) ^ (1/252)   [por dia útil, a partir de D+1]\n"
        "Pós (% do CDI):   fator ×= 1 + taxa_cdi(d) × percentual  [dia publicado, a partir de D+0]\n"
        "IPCA+:            perna prefixada diária + IPCA aplicado ao virar o mês\n"
        "Tesouro Direto:   fator = PU(d) / PU(aplicação)")

heading("B.5 Proventos", 3)
formula("provento_líquido = qtd_na_data_ex × valor_unitário × (1 − IRRF)\n\n"
        "dia de crédito = próximo pregão B3 a partir da data de PAGAMENTO")
para("Dividendos: IRRF = 0%. JCP: IRRF = 15% (até 31/12/2025) ou 17,5% (a partir de "
     "01/01/2026). O valor líquido entra em proventos_acumulados — soma ao retorno, não "
     "ao fluxo de caixa.", size=10, color=GREY)

heading("B.6 Eventos corporativos (split / grupamento / bonificação)", 3)
formula("qtd_após_evento = qtd_antes × fator\n"
        "preço_ajustado   = preço_bruto / fator   (custo total preservado)")
para("Eventos são aplicados ANTES das transações do mesmo dia. O objetivo é manter o saldo "
     "bruto contínuo na data do split — sem isso, a queda de preço do split (ex.: pela "
     "metade num 2:1) seria lida como prejuízo de ~50%.", size=10, color=GREY)

# ---- C. Carteira-brinquedo: entradas ----
heading("C. Carteira-brinquedo — dados de entrada", 2)
para("Uma única ação fictícia (AAAA3), ao longo de 5 pregões (D0 a D4, seg a sex). "
     "Eventos escolhidos para exercitar cada regra: ganho de preço, aporte adicional, "
     "dividendo e split.")
table(["Dia", "Evento", "Preço de fecho"], [
    ["D0 (seg)", "Compra de 100 ações a R$ 10,00 (aporte R$ 1.000,00)", "R$ 10,00"],
    ["D1 (ter)", "Sem operação; ação valoriza", "R$ 11,00"],
    ["D2 (qua)", "Aporte: compra +100 ações a R$ 11,00 (R$ 1.100,00)", "R$ 11,00"],
    ["D3 (qui)", "Dividendo de R$ 0,50/ação sobre 200 ações (isento) = R$ 100,00", "R$ 11,00"],
    ["D4 (sex)", "Split 2:1 (100→200... 200→400; preço cai p/ 5,50) e fecha a R$ 6,00", "R$ 6,00 (pós-split)"],
])

# ---- D. Carteira-brinquedo: cálculo dia a dia ----
heading("D. Carteira-brinquedo — cálculo dia a dia", 2)
table(["Dia", "Qtd", "Preço", "Saldo bruto", "Fluxo (aporte)", "Retorno do dia", "Acumulado"], [
    ["D0", "100", "10,00", "1.000,00", "+1.000,00", "(1000−1000)/1000 = 0,00%", "0,00%"],
    ["D1", "100", "11,00", "1.100,00", "0", "(1100−1000−0)/1000 = +10,00%", "10,00%"],
    ["D2", "200", "11,00", "2.200,00", "+1.100,00", "(2200−1100−1100)/1100 = 0,00%", "10,00%"],
    ["D3", "200", "11,00", "2.300,00", "0", "(2300−2200−0)/2200 = +4,55%", "15,00%"],
    ["D4", "400", "6,00", "2.500,00", "0", "(2500−2300−0)/2300 = +8,70%", "25,00%"],
])
para("Leitura de cada dia:", bold=True, space_after=2)
bullet("D1 — valorização pura: o saldo sobe de 1.000 para 1.100 sem dinheiro novo → +10%.")
bullet("D2 — neutralização do aporte: o saldo dobra (1.100→2.200), mas R$ 1.100 são dinheiro "
       "novo. Descontando o fluxo, o retorno do dia é 0%. O acumulado fica parado em 10%. "
       "É exatamente isto que o TWR existe para fazer.")
bullet("D3 — provento como retorno: o dividendo de R$ 100 entra no saldo (2.200→2.300) mas "
       "NÃO é fluxo de caixa. Logo conta como rentabilidade: +4,55% no dia.")
bullet("D4 — split sem ruído: a quantidade vai de 200 para 400 e o preço cai de 11,00 para "
       "5,50 (puro ajuste de escala, saldo continua 2.200). O único movimento real é 5,50→6,00, "
       "equivalente a 11,00→12,00 antes do split: +9,09% no ativo, que leva o acumulado a 25%.")
para("Saldo bruto em D4 = 400 × 6,00 + 100 (proventos acumulados) = 2.500,00.",
     size=10, color=GREY)
callout("Conferência do encadeamento: 1,00 × 1,10 × 1,00 × (2300/2200) × (2500/2300) = 1,25 → "
        "acumulado final +25,00%. Os retornos diários de D3 e D4 aparecem arredondados na "
        "tabela, mas o acumulado usa precisão cheia e fecha exatamente em 15,00% e 25,00%.")

# ---- E. Benchmarks no mesmo período ----
heading("E. Benchmarks — mesmos cinco dias (ilustração)", 2)
para("Valores ilustrativos para demonstrar o método de cada benchmark (não são dados reais "
     "do BACEN; servem para conferir a mecânica).")
para("CDI — taxa diária BACEN de 0,000400 (0,04%/dia) em D0..D4:", bold=True, space_after=2)
table(["Dia", "Índice CDI (base 100)"], [
    ["D0", "100,000000"],
    ["D1", "100,040000"],
    ["D2", "100,080016"],
    ["D3", "100,120048"],
    ["D4", "100,160096"],
])
formula("Rebaseado: (100,160096 / 100 − 1) × 100 = +0,16% na semana")
para("IBOV — se o ^BVSP fechou D0 = 120.000 e D4 = 123.000:", bold=True, space_after=2)
formula("(123.000 / 120.000 − 1) × 100 = +2,50% na semana")
para("IPCA — mês com taxa de 0,50% (mês de 31 dias), valor no 5º dia:", bold=True, space_after=2)
formula("f = (1,005) ^ (1/31) = 1,00016090\n"
        "valor(dia 5) = 100 × f^5 = 100,0805  →  +0,08% acumulado até o 5º dia")
para("Poupança segue exatamente o mesmo método mensal→diário do IPCA.", size=10, color=GREY)
para("Resultado da comparação nesta semana-exemplo: carteira +25,00% · IBOV +2,50% · "
     "CDI +0,16% · IPCA +0,08%. (Números propositalmente exagerados para a curva caber numa "
     "conferência de planilha.)", italic=True, color=GREY, size=10)

# ---- F. Checklist do consultor ----
heading("F. Roteiro de verificação para o consultor", 2)
para("Para auditar qualquer carteira real, reproduza os passos abaixo numa planilha e "
     "compare com a série exibida no gráfico:")
bullet("1. Monte o saldo bruto de cada dia útil: Σ (quantidade split-ajustada × preço de "
       "mercado) + caixa + proventos líquidos acumulados (+ renda fixa por fator, se houver).")
bullet("2. Liste os fluxos do dia (aportes − resgates). Confirme que proventos e "
       "reinvestimentos NÃO estão nessa lista.")
bullet("3. Calcule o retorno diário pela fórmula B.2 e confirme o clamp de ±50%.")
bullet("4. Encadeie (produtório) e rebaseie para 0% no primeiro dia (fórmula B.2).")
bullet("5. Para cada benchmark, reconstrua o índice (B.3) e rebaseie para 0% no mesmo dia "
       "inicial da carteira.")
bullet("6. Pontos de atenção a auditar especificamente: data de crédito dos proventos "
       "(pagamento, não data-ex), IRRF de JCP conforme a data, normalização de quantidade "
       "nos splits, e a interpolação geométrica (não linear) de IPCA/Poupança.")

import sys
out = sys.argv[1]
doc.save(out)
print("Salvo em:", out)
