#!/usr/bin/env python3
"""Gera o .docx do histórico de tentativas de paridade do gráfico de rentabilidade com o Kinvo."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x1F, 0x33, 0x5E)
GREY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1B, 0x6B, 0x3A)
RED = RGBColor(0x9B, 0x1C, 0x1C)
AMBER = RGBColor(0x8A, 0x5A, 0x00)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h

def para(text="", italic=False, color=None, size=None, bold=False, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")

def callout(text, color=GREY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run("▶  " + text)
    r.italic = True; r.font.color.rgb = color; r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(8)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val); run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============================ CAPA ============================
title = doc.add_heading("Histórico de Tentativas de Paridade com o Kinvo", level=0)
for run in title.runs: run.font.color.rgb = NAVY
para("Gráfico de Rentabilidade Diária — Análise da Carteira", bold=True, size=13)
para("MyFinance · Diário técnico para cruzamento com recomendações de consultor externo · "
     "Compilado em 15/06/2026", italic=True, color=GREY, size=10, space_after=14)
para("Objetivo deste documento: registrar TUDO o que já foi tentado para fazer nosso gráfico "
     "de rentabilidade bater com o do Kinvo — o que funcionou, o que foi descartado e por quê, "
     "e o que ainda está em aberto. A intenção é que, quando o consultor externo trouxer "
     "sugestões, você possa cruzar com esta lista e evitar refazer caminhos já percorridos. "
     "Levantado do histórico de git (≈150 commits relevantes) e do registro interno de "
     "decisões.", italic=True, color=GREY, size=10, space_after=12)

# ============================ 1. ESTADO ATUAL ============================
heading("1. Estado atual (resumo)", 1)
para("Onde chegamos, medido no ativo de teste HFOF11 (FII), comparando nossa série diária "
     "contra um capture real do Kinvo (HAR):")
bullet("As séries, quando rebaseadas ao mesmo dia inicial, batem dentro de ±0,10 ponto "
       "percentual em 2024–2025; divergência máxima em 2 anos = ~1,0 pp.")
bullet("O “penhasco-fantasma” (queda artificial de até −90% em datas de evento/provento) foi "
       "eliminado.")
bullet("A renda acumulada (proventos) bate centavo a centavo com o Kinvo (caso BBAS3: R$ 634,62).")
para("Resíduo ainda em aberto:", bold=True, space_after=2)
bullet("Sobra um resíduo de ~1 pp na ponta da série, cuja hipótese mais forte hoje é "
       "ANCORAGEM DE DATA-BASE (nós ancoramos no dia exato do aporte; o Kinvo parece ancorar "
       "no 1º dia do mês). Não é considerado erro de cálculo. Ver seção 5.")
callout("Conclusão de negócio: o gráfico está em paridade prática com o Kinvo numa comparação "
        "justa (mesma posição, mesmo dia-base). As diferenças remanescentes são de CONVENÇÃO "
        "(ancoragem) e de FONTE DE COTAÇÃO (centavos entre vendors), não de metodologia.", NAVY)

# ============================ 2. AMBIENTE DE COMPARAÇÃO ============================
heading("2. Ambiente de comparação (como medimos)", 1)
table(["Item", "Detalhe"], [
    ["Conta de teste", "qa.teste@appmyfinance.com.br, em produção (https://appmyfinance.com.br)."],
    ["Ativo-padrão", "HFOF11 (FII), 100 cotas @ R$ 73,15 em 07/06/2024. Escolhido por ter "
                     "split 10:1 e proventos mensais — exercita os casos difíceis."],
    ["Verdade-base Kinvo", "Capture real do app do Kinvo (arquivo HAR), endpoint "
                           "GetPeriodicPortfolioProfitability — série diária 2024→2026."],
    ["Harness de comparação", "Scripts internos que extraem nossa série e a do Kinvo e comparam "
                              "ponto a ponto (ver seção 7)."],
    ["Método de comparação", "Rebasear ambas ao mesmo dia inicial e comparar o passo diário — "
                             "isola erro de método de erro de ancoragem."],
])
callout("Por que HFOF11: é um FII de baixa liquidez, com split real e dividendos frequentes. "
        "Se bate nele, bate no caso difícil. Boa parte da saga abaixo nasceu dele.", GREY)

# ============================ 3. TENTATIVAS POR TEMA ============================
heading("3. Tentativas por tema", 1)
para("Cada bloco resume o que foi tentado, o resultado e o status. Legenda de status: "
     "✔ adotado · ✗revertido/descartado · ◐ parcial/em aberto.", italic=True, color=GREY, size=10)

# 3.1 TWR/MWR
heading("3.1 Metodologia de retorno (TWR / MWR)", 2)
table(["O que foi tentado", "Resultado", "Status"], [
    ["Trocar rentabilidade “só preço” por TWR acumulado (cotização, igual a fundo)",
     "TWR virou a métrica base do gráfico; alinha com a fórmula que o Kinvo documenta.", "✔"],
    ["Recalcular o TWR POR PERÍODO (não apenas filtrar o visual)",
     "Cada janela recomputa o retorno; evita distorção ao trocar de período.", "✔"],
    ["Capturar o “ganho instantâneo do dia 1” (diferença entre preço pago e preço de mercado)",
     "Alinhou o ponto inicial ao padrão Kinvo/B3 (cota = 1,00 no dia do aporte).", "✔"],
    ["Adicionar MWR (retorno ponderado pelo dinheiro) como alternativa, com toggle",
     "Disponível; em compra única buy-and-hold TWR≈MWR (esperado).", "✔"],
    ["Clamp do retorno diário (limitar saltos absurdos) — e correção do clamp inclusivo em −1",
     "Um clamp mal feito zerava o acumulado para sempre (−100% eterno). Corrigido para bounds "
     "inclusivos. Bug clássico, documentado em playbook.", "✔"],
    ["Incluir o primeiro mês na agregação mensal do TWR",
     "Corrigiu subcontagem no início da série.", "✔"],
])

# 3.2 Proventos
heading("3.2 Proventos: timing e retorno total", 2)
table(["O que foi tentado", "Resultado", "Status"], [
    ["Fazer a SÉRIE somar proventos (antes só o card somava) — retorno TOTAL",
     "Bug real: FII que caía no preço mas pagava dividendo aparecia artificialmente mal. "
     "Provento passou a entrar como retorno (não como caixa). Mudou a curva de todos os FIIs.", "✔"],
    ["Creditar o provento na DATA-EX (índice de retorno total, mais rigoroso)",
     "Implementado (jun/11). Tecnicamente o mais correto: casa o provento com a queda de preço.", "✔→◐"],
    ["Depois, creditar o provento na DATA DE PAGAMENTO (espelhar o Kinvo)",
     "Decisão posterior (jun/12) de imitar o Kinvo, que credita no pagamento. Campo bookingDay. "
     "ATENÇÃO: este ponto oscilou (ex → pagamento) — ver seção 4 e 5.", "◐"],
    ["IRRF de JCP: 15% até 31/12/2025; 17,5% a partir de 01/01/2026 (LC 224/2025)",
     "Valor líquido correto por data de pagamento.", "✔"],
    ["Elegibilidade do provento pela DATA-COM (ex-date), não pela data de pagamento",
     "Quem tinha o ativo na data-com recebe; alinha com a B3.", "✔"],
    ["Ler proventos do histórico GLOBAL por ativo (não da materialização por usuário)",
     "Eliminou o “drawdown-fantasma” para usuário novo. PortfolioProvento virou camada de "
     "override (manual vence; dismissed suprime).", "✔"],
    ["Sync robusto de proventos (dedup BRAPI, soft-delete, campo source, cron diário)",
     "Fechou gaps de renda acumulada; bateu centavo com o Kinvo.", "✔"],
    ["Segregar REINVESTIMENTO de provento do fluxo de aportes",
     "Reinvestir dividendo não infla o TWR (não é dinheiro novo).", "✔"],
])
callout("Ponto sensível nº 1 para o consultor: o crédito do provento JÁ FOI das duas formas "
        "(data-ex e data-pagamento). Hoje está na data de pagamento, para imitar o Kinvo. "
        "Se o consultor sugerir data-ex, saiba que já tentamos e é defensável — mas decidimos "
        "espelhar o Kinvo. Ver seção 5 (anomalia: a série quase não mudou ao alternar).", AMBER)

# 3.3 Splits
heading("3.3 Eventos corporativos (splits, grupamentos, bonificações)", 2)
table(["O que foi tentado", "Resultado", "Status"], [
    ["Buscar splits/grupamentos do Yahoo (a BRAPI só tem no módulo pago, que não temos)",
     "Passou a popular desdobramentos (ex.: HFOF11 10:1) que antes faltavam — a posição ficava "
     "10× errada.", "✔"],
    ["Aplicar o split à TIMELINE histórica (quantidade normalizada pós-split)",
     "Corrigiu gráfico/saldo que ficavam 10× baixos no período pré-split.", "✔"],
    ["Marcar a linha de auditoria do split como “display-only” (não conta na matemática)",
     "Sem isso, o split era contado 2× e o saldo dava spike. Regressão pega em produção.", "✔"],
    ["Dedup de eventos vindos de fontes diferentes (BRAPI bonificação + Yahoo split)",
     "Mesmo evento em datas adjacentes dobrava o fator (ex.: B3SA3 elevava ao cubo). "
     "Removidas ~399 duplicatas em prod.", "✔"],
    ["Inferir split REAL × FALSO comparando preço antes/depois do evento",
     "FALHOU: o preço da BRAPI é split-ajustado e ESCONDE o salto → removeu splits REAIS "
     "(HFOF11). Reintroduziu o penhasco. Revertido.", "✗"],
    ["Blocklist MANUAL de splits comprovadamente falsos (com evidência COTAHIST)",
     "Política final: confiar no feed por padrão; remover só falsos provados. O próprio Kinvo "
     "NÃO auto-importa eventos da B3 por má qualidade — valida nossa abordagem.", "✔"],
])
callout("Ponto sensível nº 2 para o consultor: NÃO se pode decidir se um split é real olhando "
        "o preço já ajustado (BRAPI/Yahoo recente). Só o preço CRU (COTAHIST da B3) ou uma "
        "fonte autoritativa resolve. Já erramos isso removendo splits reais. Não repetir.", AMBER)

# 3.4 Fontes de preço
heading("3.4 Fontes de preço e precedência", 2)
table(["O que foi tentado", "Resultado", "Status"], [
    ["BRAPI como fonte primária de cotação (plano PAGO — histórico completo de dividendos)",
     "Fonte principal. Corrigida a suposição antiga de “BRAPI free só tem 12 meses”.", "✔"],
    ["Yahoo Finance como fallback (preços, splits, dividendos antigos, IBOV/USD)",
     "Cobre gaps que a BRAPI não traz; backfill de 10 anos.", "✔"],
    ["B3 COTAHIST (arquivo oficial) como histórico de preço cru de anos antigos",
     "Backfill 2016→2020 (~421k cotações) e depois 2021→2026 (~838k). Resolve preço de ativos "
     "antigos — antes a falta de preço gerava rentabilidade absurda (−88%).", "✔"],
    ["Des-ajustar só as linhas COTAHIST (cru) para a escala ajustada (splitAdjustRawRows)",
     "Corrigiu o “penhasco” de ativos com preço cru + split (MXRF11, HGLG11), sem tocar "
     "nas linhas BRAPI já ajustadas. (Tentar corrigir no builder foi o caminho errado.)", "✔"],
    ["Definir PRECEDÊNCIA de fonte: B3 (COTAHIST) > BRAPI > Yahoo > … (manual vence tudo)",
     "Entregue (PR #15). A sync diária não sobrescreve mais uma cotação de fonte superior.", "✔"],
])

# 3.5 Benchmarks
heading("3.5 Benchmarks (CDI, IPCA, Poupança, IBOV)", 2)
table(["O que foi tentado", "Resultado", "Status"], [
    ["Buscar CDI/IPCA da BRAPI (atalho inicial)",
     "Substituído depois pela fonte oficial (BACEN), por precisão.", "✔→trocado"],
    ["CDI/IPCA/Poupança via BACEN SGS (séries 12 / 433 / 25), guardados no banco + cron",
     "Fonte oficial; CDI diário, IPCA e Poupança mensais. Poupança ganhou fallback próprio.", "✔"],
    ["Corrigir metodologia do CDI: compor a taxa diária (juros sobre juros), não “anualizar × 252”",
     "O método antigo distorcia o índice quando a Selic mexia. Alinhado com o pricer de renda "
     "fixa e ancorado em UTC para casar com as chaves do BACEN.", "✔"],
    ["IBOV via Yahoo (^BVSP), com backfill de 10 anos",
     "Curva diária real do índice.", "✔"],
    ["% REAL usando IPCA (inflação), em vez de “nominal − CDI”",
     "Corrigiu um cálculo que subtraía o CDI e chamava de “real” (dava −88%).", "✔"],
    ["Ingestão externa de benchmarks pré-calculados (tabela própria, prioridade máxima)",
     "Permite injetar séries oficiais já prontas; runtime lê do banco.", "✔"],
    ["Ancorar o benchmark no dia ANTERIOR (D−1) ao aporte, para fechar ~2 pp",
     "DESCARTADO de propósito — criaria assimetria (carteira em D0 × benchmark em D−1). "
     "Ver seção 4.", "✗"],
])

# 3.6 Arquitetura da série
heading("3.6 Arquitetura da série (snapshots × reconstrução)", 2)
table(["O que foi tentado", "Resultado", "Status"], [
    ["Snapshots diários persistidos (fotografias) + cauda viva, como fonte primária do gráfico",
     "Carrega rápido; o gráfico plota o snapshot e cai para reconstrução ao vivo só quando "
     "falta cobertura.", "✔"],
    ["Backfill preguiçoso de snapshots na leitura quando detecta lacuna histórica",
     "Usuário antigo/novo não fica preso a buraco de cobertura.", "✔"],
    ["Invalidação GLOBAL de snapshots após corrigir a fórmula da série",
     "Lição de processo: fix de cálculo é global → o cache de TODOS os usuários precisa ser "
     "invalidado, não só o do ativo de teste. (~3.500 snapshots removidos.)", "✔"],
    ["Invalidar snapshots em toda mutação (aporte/resgate/operação/exclusão/edição)",
     "Mantém a série coerente após qualquer mudança na carteira.", "✔"],
    ["Não plotar o dia corrente (preço intraday incompleto), igual ao Kinvo",
     "Removeu um “dente” no último ponto do gráfico.", "✔"],
    ["Eliminar o drop artificial no último dia da série",
     "Corrigido.", "✔"],
])

# 3.7 Janelas
heading("3.7 Janelas de período", 2)
table(["O que foi tentado", "Resultado", "Status"], [
    ["Alinhar filtros “últimos N meses” ao MÊS-CALENDÁRIO (dia 1º), estilo Kinvo",
     "Antes era janela rolante (hoje − N). Agora N=24 cai em 01/07/2024 (= Kinvo), não 08/06. "
     "QA de produção saiu de 3/12 para 12/12 alinhados ao Kinvo.", "✔"],
    ["Manter o gráfico de patrimônio DIÁRIO como rolante (o Kinvo também não mês-alinha o diário)",
     "Decisão consciente: só as janelas das análises são mês-alinhadas.", "✔"],
])

# ============================ 4. NÃO FAZER ============================
heading("4. Decisões de “NÃO fazer” (anti-retrabalho)", 1)
para("Esta é a seção mais importante para cruzar com o consultor. São caminhos que já foram "
     "avaliados/testados e DESCARTADOS por boa razão. Se a recomendação dele cair em algum "
     "destes, vale apresentar o motivo antes de investir tempo.", bold=True, space_after=8)
table(["Sugestão tentadora", "Por que NÃO fazemos"], [
    ["Ancorar o benchmark em D−1 para fechar os ~2 pp da ponta",
     "Nossa carteira ancora em D0 (cota = 1,00 no dia do aporte) — exatamente como o Kinvo "
     "documenta. Mudar SÓ o benchmark para D−1 criaria carteira-D0 × benchmark-D−1 "
     "(assimetria) e divergiria do método consistente. O D0 está matematicamente correto."],
    ["Decidir se um split é real comparando preço antes/depois",
     "O preço de mercado recente (BRAPI/Yahoo) já vem split-ajustado e ESCONDE o salto. "
     "Já removemos splits REAIS assim (HFOF11) e reintroduzimos o penhasco. Só preço CRU "
     "(COTAHIST) ou fonte autoritativa decide."],
    ["Auto-importar todos os eventos corporativos da B3 e confiar cegamente",
     "O próprio Kinvo NÃO faz isso — por má qualidade do dado da B3. Adotamos curadoria + "
     "blocklist manual de falsos comprovados. Confiar no feed por padrão, remover só o provado."],
    ["Assinar o módulo splitHistory da BRAPI como fonte autoritativa de split",
     "Não está no nosso plano pago (retorna 403). Avaliar só se a precisão de evento "
     "corporativo virar crítica; hoje a curadoria cobre."],
    ["Corrigir o “penhasco” de preço cru dentro do builder de patrimônio",
     "Tentado e revertido. O ajuste correto é na leitura do histórico (des-ajustar só as "
     "linhas COTAHIST), não no builder."],
    ["Perseguir paridade exata da rota de FALLBACK (carteira-historico)",
     "O fallback dá ~2,4% vs ~5,7% do snapshot no caso HFOF11 — mas ele NÃO é o que aparece no "
     "gráfico (o gráfico usa o snapshot). Diferença é de implementação (TWR ao vivo × snapshot). "
     "Baixa prioridade."],
])

# ============================ 5. EM ABERTO ============================
heading("5. Resíduos em aberto e hipóteses não fechadas", 1)
para("O que ainda não está 100% explicado — candidatos legítimos para o consultor olhar:")
table(["Item em aberto", "Situação / hipótese atual"], [
    ["Resíduo de ~1 pp na ponta da série",
     "Hipótese mais forte: ANCORAGEM de data-base (nosso aporte em 07/06 × Kinvo arredondando "
     "para o 1º do mês). Não é erro de cálculo; é convenção. Não fechado."],
    ["Anomalia: a série quase NÃO mudou ao alternar provento ex → pagamento",
     "Após trocar o crédito do provento de data-ex para data-pagamento (e trocar a fonte para "
     "B3), a série historicoTWR do HFOF11 quase não se moveu. Sugere que o ~1 pp NÃO vinha do "
     "timing do provento nem da fonte de preço (BRAPI já repassa B3). Reforça a tese de "
     "ancoragem. Próximo passo seria instrumentar o builder de patrimônio passo a passo."],
    ["Dente-de-serra mensal de ~0,85 pp em 2026",
     "Diagnóstico: é o KINVO que credita o provento só no pagamento (~2 semanas após a data-ex); "
     "aparece e some dentro do mês. Considerado artefato do Kinvo, não erro nosso."],
    ["Cobertura de proventos de FII em produção",
     "A auditoria de cobertura rodou só em dev (sem FIIs reais). Validar contra produção antes "
     "de confiar 100% no caminho para FIIs."],
    ["Caso de borda CARE11 (ajuste de split)",
     "A BRAPI às vezes guarda o preço CRU para evento recente → o des-ajuste só de COTAHIST não "
     "cobre. Resíduo conhecido e isolado."],
])
callout("Recomendação honesta: se o consultor quiser fechar o último ~1 pp, o caminho de maior "
        "retorno é investigar a ANCORAGEM de data-base (como o Kinvo escolhe o dia zero da "
        "janela), não mexer em timing de provento nem em fonte de preço — esses dois já foram "
        "exaustivamente testados e movem pouco.", NAVY)

# ============================ 6. FONTES DE DADOS ============================
heading("6. Fontes de dados: estado atual e como evoluíram", 1)
para("O usuário pediu explicitamente o histórico das fontes e suas mudanças. Estado ATUAL "
     "de cada dado e a trajetória até aqui:")
table(["Dado", "Fonte atual", "Como evoluiu"], [
    ["Preço de ações/FII/ETF",
     "B3 COTAHIST > BRAPI (paga) > Yahoo (precedência)",
     "Começou BRAPI-only → +Yahoo fallback → +COTAHIST para histórico antigo → precedência "
     "formal B3>BRAPI>Yahoo (PR #15)."],
    ["Dividendos / proventos",
     "Histórico global por ativo (banco), BRAPI paga como origem + Yahoo p/ gaps antigos",
     "BRAPI (supúnhamos free/12m) → Yahoo p/ histórico antigo → histórico GLOBAL + override "
     "manual → runtime banco-only (catálogo pré-carregado)."],
    ["Eventos corporativos (splits)",
     "Yahoo + curadoria/blocklist manual",
     "Não existiam → Yahoo (BRAPI só no módulo pago/403) → dedup cross-source → blocklist de "
     "falsos comprovados."],
    ["CDI", "BACEN SGS série 12 (diária), no banco + cron",
     "BRAPI (atalho) → BACEN; metodologia corrigida (compor diário, não anualizar×252; UTC)."],
    ["IPCA", "BACEN SGS série 433 (mensal)",
     "BRAPI → BACEN; backfill estendido para 10 anos; usado também no % real."],
    ["Poupança", "BACEN SGS série 25 (mensal)",
     "Adicionada depois; ganhou fallback próprio via economic_indexes."],
    ["IBOV", "Yahoo ^BVSP (diária)",
     "Backfill de 10 anos via Yahoo; usado como benchmark e no beta."],
    ["Benchmarks pré-calculados", "Tabela de ingestão própria (prioridade máxima)",
     "Permite injetar séries oficiais prontas; o runtime lê do banco e só busca a fonte se faltar."],
])
callout("Princípio que emergiu: runtime lê SÓ do banco (catálogo pré-carregado + rede de "
        "alarme de cobertura). Não dependemos de chamada externa no instante em que o usuário "
        "abre o gráfico — isso era frágil e reincidente.", GREY)

# ============================ 7. FERRAMENTAS DE DIAGNÓSTICO ============================
heading("7. Ferramentas de diagnóstico já existentes", 1)
para("Se o consultor quiser reproduzir/medir, este ferramental já está pronto no repositório:")
table(["Script / recurso", "Para que serve"], [
    ["scripts/dump-qa-twr.ts", "Loga como qa.teste e despeja a série historicoTWR/MWR para "
                               "comparar com o HAR do Kinvo."],
    ["scripts/debug-user-twr-mwr.ts", "Dump de carteira/transações/snapshots + saltos diários "
                                      ">|10%| persistidos (dado em DB × cálculo)."],
    ["scripts/debug-rebuild-user-twr.ts", "Reroda o builder em memória (sem persistir) e compara "
                                          "com o que está no banco."],
    ["scripts/debug-twr-step-by-step.ts", "Reproduz o cálculo dia a dia e acha o 1º ponto em que "
                                          "o acumulado zera/explode."],
    ["scripts/cleanup-stale-twr-snapshots.ts", "Varre e remove snapshots contaminados (retorno "
                                               "≤ −99%)."],
    ["scripts/audit-proventos-coverage.ts", "Gate de cobertura de proventos por símbolo/catálogo."],
    ["HAR do Kinvo", "Capture real da série do Kinvo (verdade-base para o HFOF11)."],
])

# ============================ 8. CRONOLOGIA ============================
heading("8. Cronologia de marcos (evidência em git)", 1)
para("Marcos principais, em ordem. Serve para o consultor checar datas e ler o commit "
     "correspondente se quiser o detalhe técnico.", size=10, color=GREY)
table(["Quando", "Marco", "Commit"], [
    ["Jan/2026", "Seção Análises com rentabilidade; CDI/IPCA via BRAPI (atalho inicial)", "25ce816 / 18a400a"],
    ["Fev/2026", "Fórmula vira TWR acumulado; recalcular por período", "0e0e8c9 / 978b4f1"],
    ["Fev/2026", "CDI no banco + cron; tabela de benchmarks + API de índices", "0d7954c / 1d692ec"],
    ["Mar/2026", "Snapshots diários de patrimônio; histórico global de dividendos", "7727aff / b92f927"],
    ["Abr/2026", "CDI composto (não ×252); alinhado ao pricer de RF; UTC", "c1b6650 / 616b391"],
    ["Abr/2026", "Ganho instantâneo do dia 1 no TWR (alinhamento Kinvo)", "ed2b696"],
    ["Mai/2026", "MWR + toggle TWR/MWR; clamp inclusivo −1 corrigido", "544dc1d / 76574ed"],
    ["Mai/2026", "Sync robusto de proventos; IRRF JCP 17,5%", "2cf44b6 / 15a5fac"],
    ["Mai/2026", "COTAHIST B3 2016–2020; BACEN e IBOV 10 anos", "9056e90 / ffaa041 / c391230"],
    ["Jun/2026", "Splits via Yahoo + split na timeline + auditoria display-only", "c60f130 / ad50d10 / 0c1eb68"],
    ["Jun/2026", "Janelas mês-calendário (Kinvo); gráfico alinhado (cores/escala/TWR)", "56a223b / b04ef63"],
    ["Jun/2026", "Série inclui proventos (retorno total); % real via IPCA", "fe948ac / f92e786"],
    ["Jun/2026", "Dedup de eventos cross-source; COTAHIST cliff (splitAdjustRawRows)", "—/ 01bc502"],
    ["Jun/2026", "Runtime banco-only de dividendos/eventos", "7b79192 → 18b309f"],
    ["12/Jun/2026", "PR #15: precedência B3>BRAPI>Yahoo + provento na data de pagamento", "5676ba6"],
])

# ============================ 9. CHECKLIST ============================
heading("9. Como cruzar com a recomendação do consultor", 1)
para("Roteiro rápido ao receber uma sugestão dele:")
bullet("1. A sugestão é sobre TIMING de provento (ex × pagamento)? → Já testamos os dois; "
       "hoje está no pagamento para imitar o Kinvo (seção 3.2 / 4 / 5).")
bullet("2. É sobre ANCORAGEM de benchmark/data-base? → Não mexer só no benchmark (assimetria). "
       "MAS a ancoragem de data-base da carteira é exatamente o resíduo em aberto (seção 5) — "
       "vale explorar como o Kinvo define o dia zero.")
bullet("3. É sobre SPLIT/evento corporativo? → Nunca inferir de preço ajustado; usar COTAHIST "
       "cru ou fonte autoritativa; curadoria + blocklist (seção 3.3 / 4).")
bullet("4. É sobre FONTE de preço? → Já temos precedência B3>BRAPI>Yahoo; trocar de fonte "
       "moveu pouco a série (seção 3.4 / 5).")
bullet("5. É sobre metodologia de CDI/IPCA/Poupança? → CDI é composto diário (BACEN), IPCA/"
       "Poupança são mensais interpolados (geométrico). Já corrigido (seção 3.5).")
bullet("6. É sobre a série “não bater” em valor absoluto? → Verificar primeiro se a comparação "
       "é justa (mesma posição, mesmo dia-base, snapshot × fallback) antes de assumir bug "
       "(seção 2 / 3.6).")
callout("Regra de ouro deste documento: antes de implementar uma sugestão, ache-a nas seções "
        "3 a 5. Se já está como “✔ adotado” ou “✗descartado”, você economizou um ciclo.", NAVY)

import sys
out = sys.argv[1]
doc.save(out)
print("Salvo em:", out)
